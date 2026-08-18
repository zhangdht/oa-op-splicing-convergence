# -*- coding: utf-8 -*-
"""Generate Cover_Letter.docx for the OA-OP splicing manuscript."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# base style
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(11)

def p(text, bold=False, space_after=6, align=None):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    return para

# ---- header ----
p("2 August 2026", space_after=2)
p("Editor-in-Chief", bold=True, space_after=2)
p("RMD Open", space_after=2)
p("", space_after=2)

p("Dear Editor,", space_after=10)

p("Re: Submission of manuscript \u201cCompartment-specific convergence of the RNA splicing "
  "programme in osteoarthritis and osteoporosis\u201d", bold=True, space_after=10)

p("We enclose our manuscript for consideration as an Original research article in "
  "RMD Open.", space_after=8)

p("Osteoarthritis (OA) and osteoporosis (OP) are conventionally analysed as opposite diseases, yet "
  "their molecular relationship remains disputed. Most transcriptomic comparisons pool heterogeneous "
  "tissues under a single disease label. In this study we re-examine seven public cohorts (303 samples "
  "spanning OA cartilage, OP circulating monocytes and OP bone-marrow mesenchymal stem cells) and show "
  "that the transcriptomic link between OA and OP is compartment-specific rather than disease-specific: a "
  "coordinated down-regulation of the RNA splicing programme unites articular cartilage and bone-marrow "
  "stroma (both mesenchymal), whereas the haematopoietic compartment follows a separate, partly opposite "
  "trajectory. Pooling OP tissues cancels this signal, which we propose explains why previous searches for "
  "a shared OA-OP hub have been inconsistent.", space_after=8)

# --- the explicit sentence the user required ---
p("This is a re-analysis that resolves apparent contradictions in Chen 2024 / Wang 2025 / Jiang 2026 by "
  "tissue-compartment stratification, not a replication attempt.", bold=True, space_after=8)

p("We therefore do not claim to replicate those studies; rather, we show that their divergent hub-gene "
  "lists are reconciled once tissues are assigned to the correct developmental compartment, and we "
  "contextualise the recent nomination of SON as the causal OP-OA hub as a passenger of a mesenchymal "
  "splicing module rather than its driver. This framing should pre-empt any editorial reading of the work "
  "as a failed replication.", space_after=8)

p("RMD Open reaches rheumatologists, musculoskeletal epidemiologists and translational researchers, the "
  "exact audience for a compartment-aware re-interpretation of the OA-OP molecular relationship. By "
  "reconciling apparently contradictory hub-gene reports using only openly available data that the "
  "community can independently re-run, the study offers a reusable analytical template for musculoskeletal "
  "comorbidity research.", space_after=8)

p("All authors have read and approved the final manuscript. All analyses used only publicly available, "
  "de-identified datasets. The authors declare no competing interests. Analysis code and intermediate "
  "results are deposited in a public repository (archived at Zenodo) and are publicly accessible. We "
  "confirm that this work is original, has not been published elsewhere, and is not under consideration by "
  "another journal.", space_after=8)

p("We believe this work will interest your readership and look forward to your assessment.", space_after=8)

p("Sincerely,", space_after=2)
p("Zhang Dong, MD, Chief Physician", space_after=2)
p("Corresponding author", space_after=2)
p("Email: zhangdht@126.com", space_after=2)
p("ORCID: 0009-0006-9000-273X", space_after=2)

out = r"D:/bioinfo05/06_paper/Cover_Letter.docx"
doc.save(out)
print("Saved:", out)
