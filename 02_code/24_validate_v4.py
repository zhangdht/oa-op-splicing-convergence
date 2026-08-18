"""v4 专项校验：作者上标 + 前轮 17 项回归。"""
import os, re, zipfile, csv
from docx import Document
from docx.oxml.ns import qn

ROOT = "D:/bioinfo05"
DOCX = os.path.join(ROOT, "06_paper", "Manuscript_final_v4.docx")

assert zipfile.ZipFile(DOCX).testzip() is None
d = Document(DOCX)
for t in d.tables:
    assert len(t.columns) > 0
full = "\n".join(p.text for p in d.paragraphs)
checks = []

# 1. 作者上标
p = d.paragraphs[2]
sup = [r.text for r in p.runs
       if r._r.find(qn('w:rPr') + '/' + qn('w:vertAlign')) is not None
       and r._r.find(qn('w:rPr') + '/' + qn('w:vertAlign')).get(qn('w:val')) == 'superscript']
checks.append(("Author line has superscript 1/2/* markers",
               set(sup) >= {"1", "2", "*"}))
checks.append(("Author names present (Wei/Li/Chen/Zhang)",
               all(x in full for x in ["Wei Yuwei", "Li Peng", "Chen Kai", "Zhang Dong"])))

# 2. 引用连续 1-26
nums = set()
for m in re.finditer(r"\[([0-9,\s]+)\]", full):
    for part in m.group(1).split(","):
        part = part.strip()
        if part:
            nums.add(int(part))
checks.append(("Citations contiguous 1..26", nums == set(range(1, 27))))

# 3. Declarations 填实
checks.append(("No 'to be completed' left", "to be completed" not in full.lower()))
checks.append(("Funding embedded (2025HK044)", "2025HK044" in full))
checks.append(("Corresponding author line present", "Corresponding author: Zhang Dong" in full))

# 4. Abstract FDR
checks.append(("Abstract 'all FDR > 0.99'", "maximum Jaccard 0.032; all FDR > 0.99" in full))

# 5. Table 3 符号
t3 = d.tables[3]
rows = {r.cells[0].text.strip(): [c.text.strip() for c in r.cells] for r in t3.rows[1:]}
checks.append(("T3 REGULATION OF RNA SPLICING Z OA -4.86",
               rows.get("REGULATION OF RNA SPLICING", [""])[2].startswith("-4.86")))
checks.append(("T3 POSITIVE REGULATION Z OA -4.54",
               rows.get("POSITIVE REGULATION OF RNA SPLICING", [""])[2].startswith("-4.54")))
checks.append(("T3 ALTERNATIVE... Z OA == -3.40",
               rows.get("ALTERNATIVE MRNA SPLICING VIA SPLICEOSOME", [""])[2] == "-3.40"))

# 6. Table 2a vs
t2 = d.tables[1]
t2text = " | ".join(" | ".join(c.text.strip() for c in r.cells) for r in t2.rows)
checks.append(("Table 2a 'GSE57218 vs GSE169077'", "GSE57218 vs GSE169077" in t2text))

# 7. 11/49 脚注统一
checks.append(("Table 4 footnote unified to 11/49 concordantly",
               "11 of 49 (22%) were concordantly down-regulated" in full))
checks.append(("No stale '22 of 49 (45%)' in footnote",
               "22 of 49 (45%)" not in full))

# 8. 公开仓库
checks.append(("Methods repo public (Zenodo)",
               "Zenodo" in full and "public GitHub repository" in full))

# 9. Figures 1-6
figcaps = sum(1 for p in d.paragraphs if re.match(r'^Figure \d+\.', p.text.strip()))
checks.append(("6 Figure captions", figcaps == 6))

# 10. Ref [3] 干净（参考文献按 "3." 编号）
refs3 = [p.text for p in d.paragraphs if re.match(r'^3\.', p.text.strip())]
checks.append(("Ref [3] Peng 2026 clean, correct DOI",
               refs3 and "10.1016/j.joca.2026.05.013" in refs3[0]
               and "42218990" not in refs3[0]))  # PMID 不应内联

ok = True
print("=== v4 validation ===")
for name, res in checks:
    print(("PASS" if res else "FAIL"), "-", name)
    ok = ok and res
print("\nALL GOOD" if ok else "\nSOME CHECKS FAILED")
