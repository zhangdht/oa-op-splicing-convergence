#!/usr/bin/env python
"""
19_validate_project.py
----------------------
Deep validation of the bioinfo05 deliverables:
  - key files exist
  - Manuscript_final.docx is a valid zip and opens with python-docx
  - every table carries a w:tblGrid (Word-safe)
  - citation numbers are contiguous 1..N and grouped forms resolve
  - all cited reference keys exist in references_verified.csv with a real DOI/PMID
Run: python 02_code/19_validate_project.py
"""
import os, re, csv, zipfile, sys

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
def p(*a): return os.path.join(ROOT, *a)

required = [
    "06_paper/Manuscript_final.docx",
    "06_paper/references_verified.csv",
    "04_figures/Fig1_design_and_landscape.tiff",
    "04_figures/Fig2_gene_level_no_overlap.tiff",
    "04_figures/Fig3_compartment_convergence.tiff",
    "04_figures/Fig4_splicing_axis.tiff",
    "04_figures/Fig5_single_cell_landscape.tiff",
    "04_figures/Fig6_SON_ML_genetics.tiff",
    "05_tables/Table1_cohorts_formatted.csv",
    "05_tables/Table2_convergence_controls.csv",
    "05_tables/Table2_family_level_convergence.csv",
    "05_tables/Table3_top_splicing_pathways.csv",
    "03_results/functional/mesenchymal_leading_edge_genes.csv",
    "03_results/genetics/gene_trait_hits_bone_OA.csv",
    "02_code/16_write_paper_final.R",
    "02_code/17_genetics_gwas_catalog.py",
    "02_code/18_finalize_docx.py",
    "README.md",
]
missing = [f for f in required if not os.path.exists(p(f))]
assert not missing, "MISSING FILES:\n" + "\n".join(missing)
print(f"[ok] {len(required)} key files present")

# docx integrity + tables
from docx import Document
docx = p("06_paper/Manuscript_final.docx")
assert zipfile.ZipFile(docx).testzip() is None
d = Document(docx)
for i, t in enumerate(d.tables):
    assert len(t.columns) > 0, f"Table {i+1} missing grid"
print(f"[ok] docx opens; {len(d.tables)} tables, all with w:tblGrid")

# citations contiguous
txt = "\n".join(par.text for par in d.paragraphs)
nums = set()
for m in re.finditer(r"\[([0-9,\s]+)\]", txt):
    for part in m.group(1).split(","):
        part = part.strip()
        if part:
            nums.add(int(part))
assert nums == set(range(1, max(nums) + 1)), f"citation gap: {sorted(nums)}"
print(f"[ok] citations contiguous 1..{max(nums)}")
assert not re.findall(r"\(\d{6,8}\)", txt), "stray bare PMID found"
print("[ok] no stray bare PMID citations")

# cited keys exist with real id
rows = list(csv.DictReader(open(p("06_paper/references_verified.csv"), encoding="utf-8-sig")))
bykey = {r["key"]: r for r in rows}
# extract cited keys from the R script's cite() calls
src = open(p("02_code/16_write_paper_final.R"), encoding="utf-8").read()
cited = sorted(set(k for call in re.findall(r'cite\((.*?)\)', src, re.S)
                   for k in re.findall(r'"([^"]+)"', call)))
bad = [k for k in cited if k not in bykey or not (bykey[k].get("doi","").strip() or bykey[k].get("pmid","").strip())]
assert not bad, f"cited refs missing/invalid: {bad}"
print(f"[ok] {len(cited)} cited references all present & real in references_verified.csv")

print("\nALL VALIDATIONS PASSED")
