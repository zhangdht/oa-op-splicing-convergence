import re, zipfile
from docx import Document
DOCX = r"D:/bioinfo05/06_paper/Manuscript_final_v3.docx"
assert zipfile.ZipFile(DOCX).testzip() is None
d = Document(DOCX)
for t in d.tables:
    assert len(t.columns) > 0

full = "\n".join(p.text for p in d.paragraphs)
checks = []

# citations contiguous
nums = set()
for m in re.finditer(r"\[([0-9,\s]+)\]", full):
    for part in m.group(1).split(","):
        part = part.strip()
        if part:
            nums.add(int(part))
checks.append(("Citations contiguous 1..%d" % max(nums), nums == set(range(1, max(nums)+1))))
checks.append(("No stray bare PMID like (12345678)", not re.findall(r"\(\d{6,8}\)", full)))

# [1] Table 3 (splicing) all Z OA negative + Direction Down in both
t3 = d.tables[3]
rows = {r.cells[0].text.strip(): [c.text.strip() for c in r.cells] for r in t3.rows[1:]}
allneg = all(r[2].startswith("-") and r[6] == "Down in both" for r in rows.values())
checks.append(("[1] Table 3 all Z OA negative & Down in both", allneg))
checks.append(("[1] ALTERNATIVE MRNA SPLICING VIA SPLICEOSOME == -3.40", rows.get("ALTERNATIVE MRNA SPLICING VIA SPLICEOSOME", [""]*7)[2] == "-3.40"))
checks.append(("[1] REGULATION OF RNA SPLICING == -4.86", rows.get("REGULATION OF RNA SPLICING", [""]*7)[2] == "-4.86"))
checks.append(("[1] POSITIVE REGULATION ... == -4.54", rows.get("POSITIVE REGULATION OF RNA SPLICING", [""]*7)[2] == "-4.54"))

# [2] Table 2a vs
t2 = d.tables[1]
t2text = " | ".join(" | ".join(c.text.strip() for c in r.cells) for r in t2.rows)
checks.append(("[2] Table 2a 'GSE57218 vs GSE169077'", "GSE57218 vs GSE169077" in t2text))

# [3] Table 4 footnote 11/49 (not 22/49) + matches body
t4 = d.tables[4]
foot = " | ".join(c.text.strip() for c in t4.rows[0].cells)  # title+footnote are paragraphs, not in table; grab from full text
# footnote text is in paragraph after the table title; search in full
checks.append(("[3] Table 4 footnote says '11 of 49 (22%)'", "11 of 49 (22%)" in full))
checks.append(("[3] Table 4 footnote NO '22 of 49 (45%)'", "22 of 49 (45%)" not in full))

# [4] Figure 6 present: 6 captions + 6 drawings
from docx.oxml.ns import qn
ndraw = sum(len(p._p.findall('.//'+qn('w:drawing'))) for p in d.paragraphs)
ncaps = sum(1 for p in d.paragraphs if re.match(r'^Figure \d+\.', p.text.strip()))
checks.append(("[4] 6 figure drawings present", ndraw == 6))
checks.append(("[4] 6 'Figure N.' captions", ncaps == 6))
checks.append(("[4] Figure 6 caption present", any(p.text.strip().startswith("Figure 6.") for p in d.paragraphs)))

# [5] Abstract FDR
checks.append(("[5] Abstract 'all FDR > 0.99'", "maximum Jaccard 0.032; all FDR > 0.99" in full))

# [7] figure numbering (no literal 'image' refs)
checks.append(("[7] No literal 'image' figure refs", "image" not in full.lower()))

# [9] Methods repo public
checks.append(("[9] Methods 'deposited in a public' (not 'on request')", "deposited in a public" in full and "available from the corresponding author on request" not in full))

# [6] references [3] peng (DOI, advance) / [8] jiang (vol 27)
refs = []
grab = False
for p in d.paragraphs:
    t = p.text.strip()
    if t == "References":
        grab = True
    if grab:
        m = re.match(r'^(\d+)\.\s', t)
        if m:
            refs.append((int(m.group(1)), t))
ref3 = next((t for n, t in refs if n == 3), "")
ref8 = next((t for n, t in refs if n == 8), "")
checks.append(("[6] [3] Peng 2026 has doi (advance, no vol flagged by DOI)", "doi:10.1016/j.joca.2026.05.013" in ref3))
checks.append(("[6] [8] Jiang 2026 has volume 27", "International journal of molecular sciences. 2026. doi:10.3390/ijms27114905" in ref8))

print("=== VALIDATION (all 10 points) ===")
ok = True
for name, res in checks:
    print(("PASS" if res else "FAIL"), "-", name)
    ok = ok and res
print("\nRESULT:", "ALL PASS" if ok else "SOME FAILED")
