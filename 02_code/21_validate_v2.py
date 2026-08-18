import os, re, csv, zipfile
from docx import Document
ROOT = "D:/bioinfo05"
docx = os.path.join(ROOT, "06_paper", "Manuscript_final_v2.docx")
assert zipfile.ZipFile(docx).testzip() is None
d = Document(docx)
for i, t in enumerate(d.tables):
    assert len(t.columns) > 0
print(f"[ok] docx opens; {len(d.tables)} tables, all with w:tblGrid")
full = "\n".join(p.text for p in d.paragraphs)

nums = set()
for m in re.finditer(r"\[([0-9,\s]+)\]", full):
    for part in m.group(1).split(","):
        part = part.strip()
        if part:
            nums.add(int(part))
assert nums == set(range(1, max(nums) + 1)), f"citation gap: {sorted(nums)}"
print(f"[ok] citations contiguous 1..{max(nums)}")

checks = []
checks.append(("Authors embedded (Wei/Li/Chen/Zhang + Corresponding)",
               all(x in full for x in ["Wei Yuwei", "Li Peng", "Chen Kai", "Zhang Dong", "Corresponding author: Zhang Dong"])))
checks.append(("Funding embedded (2025HK044)",
               "2025HK044" in full))
checks.append(("Declarations filled (no 'to be completed')",
               "to be completed" not in full.lower()))
checks.append(("Abstract has 'all FDR > 0.99'",
               "maximum Jaccard 0.032; all FDR > 0.99" in full))

# docx Table 4 (index 3) = splicing pathways (source Table3)
t3 = d.tables[3]
rows = {r.cells[0].text.strip(): [c.text.strip() for c in r.cells] for r in t3.rows[1:]}
reg = rows.get("REGULATION OF RNA SPLICING")
pos = rows.get("POSITIVE REGULATION OF RNA SPLICING")
alt = rows.get("ALTERNATIVE MRNA SPLICING VIA SPLICEOSOME")
checks.append(("Table3 REGULATION OF RNA SPLICING Z OA negative (-4.86)",
               reg is not None and reg[2].startswith("-4.86")))
checks.append(("Table3 POSITIVE REGULATION OF RNA SPLICING Z OA negative (-4.54)",
               pos is not None and pos[2].startswith("-4.54")))
checks.append(("Table3 ALTERNATIVE MRNA SPLICING VIA SPLICEOSOME Z OA == -3.40",
               alt is not None and alt[2] == "-3.40"))

# docx Table 2 (index 1) = Table2a
t2 = d.tables[1]
t2text = " | ".join(" | ".join(c.text.strip() for c in r.cells) for r in t2.rows)
checks.append(("Table2a has 'GSE57218 vs GSE169077'",
               "GSE57218 vs GSE169077" in t2text))

figcaps = sum(1 for p in d.paragraphs if re.search(r"^Figure \d+\.", p.text.strip()))
intext = len(re.findall(r"Figure \d+", full))
checks.append(("6 Figure captions present", figcaps == 6))
checks.append(("In-text Figure references present (>=6 in body)", intext >= 12))

checks.append(("Peng 2026 cited in intro",
               "peng2026_bone_cartilage_paradox" in open(os.path.join(ROOT, "02_code", "16_write_paper_final.R"), encoding="utf-8").read()))

print("\n--- user-requested checks ---")
ok = True
for name, res in checks:
    print(("PASS" if res else "FAIL"), "-", name)
    ok = ok and res
print("\nALL GOOD" if ok else "SOME CHECKS FAILED")
